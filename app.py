# -*- coding: utf-8 -*-

import streamlit as st
import pandas as pd
import vertexai
from vertexai.preview.generative_models import GenerativeModel
import PyPDF2
import docx
import re
import io

# --- CONFIGURACIÓN DE LA PÁGINA DE STREAMLIT ---
st.set_page_config(
    page_title="Generador y Auditor de Ítems con IA (Vertex AI)",
    page_icon="🧠",
    layout="wide"
)

# --- VARIABLES DE ENTORNO ---
# Estas variables se configuran en el entorno de despliegue (Cloud Run).
GCP_PROJECT_ID = os.environ.get("GCP_PROJECT_ID")
GCP_LOCATION = os.environ.get("GCP_LOCATION")

# --- INICIALIZACIÓN DE VERTEX AI ---
# Se ejecuta una sola vez y usa la autenticación del entorno de Cloud Run.
if 'vertex_initialized' not in st.session_state:
    try:
        if GCP_PROJECT_ID and GCP_LOCATION:
            vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
            st.session_state.vertex_initialized = True
            st.sidebar.success("✅ Conectado a Vertex AI.")
        else:
            st.sidebar.error("Variables de entorno GCP_PROJECT_ID o GCP_LOCATION no configuradas.")
            st.session_state.vertex_initialized = False
    except Exception as e:
        st.sidebar.error(f"Error al inicializar Vertex AI: {e}")
        st.session_state.vertex_initialized = False


# --- Funciones de Lectura de Archivos ---
@st.cache_data
def leer_excel_cargado(uploaded_file):
    if uploaded_file is not None:
        try:
            df = pd.read_excel(uploaded_file)
            st.sidebar.success(f"Archivo Excel '{uploaded_file.name}' cargado exitosamente.")
            return df
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el archivo Excel: {e}")
            return None
    return None

@st.cache_data
def leer_pdf_cargado(uploaded_file):
    if uploaded_file is not None:
        try:
            texto_pdf = ""
            reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            for page_num in range(len(reader.pages)):
                texto_pdf += reader.pages[page_num].extract_text()
            st.sidebar.success(f"Archivo PDF '{uploaded_file.name}' leído exitosamente.")
            return texto_pdf
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el archivo PDF: {e}")
            return ""
    return ""

# --- Función para obtener la descripción de la taxonomía de Bloom ---
def get_descripcion_bloom(proceso_cognitivo_elegido):
    descripcion_bloom_map = {
        "RECORDAR": "Recuperar información relevante desde la memoria de largo plazo.",
        "COMPRENDER": "Construir significado a partir de información mediante interpretación, resumen, explicación u otras tareas.",
        "APLICAR": "Usar procedimientos en situaciones conocidas o nuevas.",
        "ANALIZAR": "Descomponer información y examinar relaciones entre partes.",
        "EVALUAR": "Emitir juicios basados en criterios para valorar ideas o soluciones.",
        "CREAR": "Generar nuevas ideas, productos o formas de reorganizar información."
    }
    return descripcion_bloom_map.get(str(proceso_cognitivo_elegido).upper(), "Descripción no disponible.")

# --- Función para generar texto con Vertex AI ---
def generar_texto_con_llm(model_name, prompt):
    """
    Genera texto usando un modelo de Gemini a través de la infraestructura de Vertex AI.
    """
    if 'vertex_initialized' not in st.session_state or not st.session_state.vertex_initialized:
        st.error("Vertex AI no está inicializado. Verifica la configuración.")
        return None
    
    try:
        modelo = GenerativeModel(model_name)
        response = modelo.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error al llamar al modelo '{model_name}' en Vertex AI: {e}")
        return None

# --- Función para auditar el ítem generado ---
def auditar_item_con_llm(model_name, item_generado, grado, area, asignatura, estacion,
                         proceso_cognitivo, nanohabilidad, microhabilidad,
                         competencia_nanohabilidad, contexto_educativo, manual_reglas_texto="", descripcion_bloom="", grafico_necesario="", descripcion_grafico="", prompt_auditor_adicional=""):
    auditoria_prompt = f"""
    Eres un experto en validación de ítems educativos, especializado en pruebas tipo ICFES y las directrices del equipo IMPROVE.
    Tu tarea es AUDITAR RIGUROSAMENTE el siguiente ítem generado por un modelo de lenguaje.
    Debes verificar que el ítem cumpla con TODOS los siguientes criterios.

    --- CRITERIOS DE AUDITORÍA ---
    1.  **Formato del Enunciado:** ¿El enunciado está formulado como pregunta clara y directa?
    2.  **Número de Opciones:** ¿Hay exactamente 4 opciones (A, B, C, D)?
    3.  **Respuesta Correcta Indicada:** ¿La sección 'RESPUESTA CORRECTA:' está claramente indicada?
    4.  **Diseño de Justificaciones:** ¿Las justificaciones siguen el formato requerido (explicación para la correcta, análisis de error para las incorrectas)?
    5.  **Estilo y Restricciones:** ¿No se usan negaciones mal redactadas, nombres reales, marcas, etc.?
    6.  **Alineación del Contenido:** ¿El ítem está alineado EXCLUSIVAMENTE con los siguientes elementos?
        * Grado: {grado}
        * Área: {area}
        * Asignatura: {asignatura}
        * Estación o unidad temática: {estacion}
        * Proceso Cognitivo (Taxonomía de Bloom): {proceso_cognitivo} (descripción: "{descripcion_bloom}")
        * Nanohabilidad: {nanohabilidad}
        * Microhabilidad: {microhabilidad}
        * Competencia: {competencia_nanohabilidad}
        * Nivel educativo: {contexto_educativo}
    7.  **Gráfico (si aplica):** Si el ítem indica que requiere un gráfico, ¿la descripción es clara?
        * Gráfico Necesario: {grafico_necesario}
        * Descripción del Gráfico: {descripcion_grafico if grafico_necesario == 'SÍ' else 'N/A'}

    --- MANUAL DE REGLAS ADICIONAL ---
    {manual_reglas_texto}
    -----------------------------------

    --- INSTRUCCIONES ADICIONALES PARA LA AUDITORÍA ---
    {prompt_auditor_adicional if prompt_auditor_adicional else "No se proporcionaron instrucciones adicionales."}
    ---------------------------------------------------

    ÍTEM A AUDITAR:
    --------------------
    {item_generado}
    --------------------

    Devuelve tu auditoría con este formato estructurado:

    VALIDACIÓN DE CRITERIOS:
    - Formato del Enunciado: [✅ / ❌] + Comentario (si ❌)
    - Número de Opciones (4): [✅ / ❌]
    - Respuesta Correcta Indicada: [✅ / ❌]
    - Diseño de Justificaciones: [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)
    - Estilo y Restricciones: [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)
    - Alineación del Contenido: [✅ / ❌] + Comentario (si ❌)
    - Gráfico (si aplica): [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)

    DICTAMEN FINAL:
    [✅ CUMPLE TOTALMENTE / ⚠️ CUMPLE PARCIALMENTE / ❌ RECHAZADO]

    OBSERVACIONES FINALES:
    [Explica de forma concisa qué aspectos necesitan mejora, si el dictamen no es ✅.]
    """
    return generar_texto_con_llm(model_name, auditoria_prompt), auditoria_prompt

# --- Función principal para generar y auditar preguntas ---
def generar_pregunta_con_seleccion(gen_model_name, audit_model_name,
                                   fila_datos, criterios_generacion, manual_reglas_texto="",
                                   informacion_adicional_usuario="",
                                   prompt_bloom_adicional="", prompt_construccion_adicional="", prompt_especifico_adicional="",
                                   prompt_auditor_adicional="",
                                   contexto_general_estacion=""):
    tipo_pregunta = criterios_generacion.get("tipo_pregunta", "opción múltiple con 4 opciones")
    dificultad = criterios_generacion.get("dificultad", "media")
    contexto_educativo = criterios_generacion.get("contexto_educativo", "general")
    formato_justificacion = criterios_generacion.get("formato_justificacion", """
        • Justificación correcta: debe explicar el razonamiento o proceso cognitivo (NO por descarte).
        • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
    """)
    
    grado_elegido = fila_datos.get('GRADO', 'no especificado')
    area_elegida = fila_datos.get('ÁREA', 'no especificada')
    asignatura_elegida = fila_datos.get('ASIGNATURA', 'no especificada')
    estacion_elegida = fila_datos.get('ESTACIÓN', 'no especificada')
    proceso_cognitivo_elegido = fila_datos.get('PROCESO COGNITIVO', 'no especificado')
    nanohabilidad_elegida = fila_datos.get('NANOHABILIDAD', 'no especificada')
    microhabilidad_elegida = fila_datos.get('MICROHABILIDAD', 'no especificada')
    competencia_nanohabilidad_elegida = fila_datos.get('COMPETENCIA NANOHABILIDAD', 'no especificada')

    descripcion_bloom = get_descripcion_bloom(proceso_cognitivo_elegido)

    current_item_text = ""
    auditoria_status = "❌ RECHAZADO"
    audit_observations = ""
    max_refinement_attempts = 5
    attempt = 0
    grafico_necesario = "NO"
    descripcion_grafico = ""
    item_final_data = None
    full_generation_prompt = ""
    full_auditor_prompt = ""

    classification_details = {
        "Grado": grado_elegido, "Área": area_elegida, "Asignatura": asignatura_elegida,
        "Estación": estacion_elegida, "Proceso Cognitivo": proceso_cognitivo_elegido,
        "Nanohabilidad": nanohabilidad_elegida, "Microhabilidad": microhabilidad_elegida,
        "Competencia Nanohabilidad": competencia_nanohabilidad_elegida
    }

    while auditoria_status != "✅ CUMPLE TOTALMENTE" and attempt < max_refinement_attempts:
        attempt += 1
        
        prompt_content_for_llm = f"""
        Eres un diseñador experto en ítems de evaluación educativa, especializado en pruebas tipo ICFES.
        Tu tarea es construir un ítem de {tipo_pregunta} con una única respuesta correcta.

        --- CONTEXTO Y PARÁMETROS DEL ÍTEM ---
        - Grado: {grado_elegido}
        - Área: {area_elegida}
        - Asignatura: {asignatura_elegida}
        - Estación o unidad temática: {estacion_elegida}
        - Proceso cognitivo (Taxonomía de Bloom): {proceso_cognitivo_elegido}
        - Descripción del proceso cognitivo: "{descripcion_bloom}"
        
        --- PROMPT ADICIONAL: TAXONOMÍA DE BLOOM / PROCESOS COGNITIVOS ---
        {prompt_bloom_adicional if prompt_bloom_adicional else "No se proporcionaron prompts adicionales."}
        ------------------------------------------------------------------

        - Nanohabilidad (foco principal): {nanohabilidad_elegida}
        - Nivel educativo: {contexto_educativo}
        - Dificultad deseada: {dificultad}

        --- CONTEXTO GENERAL DE LA ESTACIÓN (si aplica) ---
        {f"Considera este contexto general: {contexto_general_estacion}" if contexto_general_estacion else "Genera un contexto individual para este ítem."}
        ----------------------------------------------------

        --- INSTRUCCIONES PARA LA CONSTRUCCIÓN ---
        CONTEXTO: Incluye una situación relevante para el grado y área.
        ENUNCIADO: Formula una pregunta clara y directa. Si usas negaciones, resáltalas en MAYÚSCULAS Y NEGRITA.
        OPCIONES: Escribe exactamente cuatro opciones (A, B, C, D). Solo una correcta. Los distractores deben ser creíbles.
        JUSTIFICACIONES: {formato_justificacion}

        --- PROMPT ADICIONAL: REGLAS GENERALES DE CONSTRUCCIÓN ---
        {prompt_construccion_adicional if prompt_construccion_adicional else "No se proporcionaron prompts adicionales."}
        ---------------------------------------------------------

        --- REGLAS ADICIONALES DEL MANUAL ---
        Aplica estrictamente las directrices del siguiente manual:
        {manual_reglas_texto}
        ----------------------------------------------------

        --- INFORMACIÓN ADICIONAL DEL USUARIO ---
        {informacion_adicional_usuario if informacion_adicional_usuario else "No se proporcionó información adicional."}
        ---------------------------------------------------------------------------
        
        --- PROMPT ADICIONAL: COSAS ESPECÍFICAS A TENER EN CUENTA ---
        {prompt_especifico_adicional if prompt_especifico_adicional else "No se proporcionaron prompts adicionales."}
        ----------------------------------------------------------

        --- DATO CLAVE PARA LA CONSTRUCCIÓN ---
        Basa el ítem en la siguiente nanohabilidad: "{nanohabilidad_elegida}"

        --- INSTRUCCIONES DE SALIDA PARA GRÁFICO ---
        Después de las justificaciones, incluye:
        GRAFICO_NECESARIO: [SÍ/NO]
        DESCRIPCION_GRAFICO: [Si es SÍ, descripción detallada. Si es NO, escribe N/A.]

        --- FORMATO ESPERADO DE SALIDA ---
        PREGUNTA: [Enunciado]
        A. [Opción A]
        B. [Opción B]
        C. [Opción C]
        D. [Opción D]
        RESPUESTA CORRECTA: [Letra]
        JUSTIFICACIONES:
        A. [Justificación A]
        B. [Justificación B]
        C. [Justificación C]
        D. [Justificación D]
        GRAFICO_NECESARIO: [SÍ/NO]
        DESCRIPCION_GRAFICO: [Descripción o N/A]
        """
        
        if attempt > 1:
            prompt_content_for_llm += f"""
            --- RETROALIMENTACIÓN DE AUDITORÍA PARA REFINAMIENTO ---
            El ítem anterior no cumplió los criterios. Revisa estas observaciones y mejora el ítem:
            Observaciones del Auditor: {audit_observations}
            --- ÍTEM ANTERIOR A REFINAR ---
            {current_item_text}
            -------------------------------
            """
        
        full_generation_prompt = prompt_content_for_llm

        try:
            with st.spinner(f"Generando contenido con IA ({gen_model_name}, Intento {attempt})..."):
                full_llm_response = generar_texto_con_llm(gen_model_name, prompt_content_for_llm)
                
                if full_llm_response is None:
                    auditoria_status = "❌ RECHAZADO (Error de Generación)"
                    audit_observations = "El modelo de generación no pudo producir una respuesta."
                    break
                
                item_and_graphic_match = re.search(r"(PREGUNTA:.*?)(GRAFICO_NECESARIO:\s*(SÍ|NO).*?DESCRIPCION_GRAFICO:.*)", full_llm_response, re.DOTALL)
                
                if item_and_graphic_match:
                    current_item_text = item_and_graphic_match.group(1).strip()
                    grafico_info_block = item_and_graphic_match.group(2).strip()
                    
                    grafico_necesario_match = re.search(r"GRAFICO_NECESARIO:\s*(SÍ|NO)", grafico_info_block)
                    grafico_necesario = grafico_necesario_match.group(1).strip() if grafico_necesario_match else "NO"

                    descripcion_grafico_match = re.search(r"DESCRIPCION_GRAFICO:\s*(.*)", grafico_info_block, re.DOTALL)
                    descripcion_grafico = descripcion_grafico_match.group(1).strip() if descripcion_grafico_match else ""
                    if descripcion_grafico.upper() == 'N/A':
                        descripcion_grafico = ""
                else:
                    current_item_text = full_llm_response
                    grafico_necesario = "NO"
                    descripcion_grafico = ""
                    st.warning("No se pudo parsear el formato de gráfico. Asumiendo que no se requiere.")

            with st.spinner(f"Auditando ítem ({audit_model_name}, Intento {attempt})..."):
                auditoria_resultado, full_auditor_prompt = auditar_item_con_llm(
                    audit_model_name,
                    item_generado=current_item_text,
                    grado=grado_elegido, area=area_elegida, asignatura=asignatura_seleccionada, estacion=estacion_elegida,
                    proceso_cognitivo=proceso_cognitivo_seleccionado, nanohabilidad=nanohabilidad_seleccionada,
                    microhabilidad=microhabilidad_elegida, competencia_nanohabilidad=competencia_nanohabilidad_elegida,
                    contexto_educativo=contexto_educativo, manual_reglas_texto=manual_reglas_texto,
                    descripcion_bloom=descripcion_bloom,
                    grafico_necesario=grafico_necesario,
                    descripcion_grafico=descripcion_grafico,
                    prompt_auditor_adicional=prompt_auditor_adicional
                )
                if auditoria_resultado is None:
                    auditoria_status = "❌ RECHAZADO (Error de Auditoría)"
                    audit_observations = "El modelo de auditoría no pudo producir una respuesta."
                    break

            dictamen_final_match = re.search(r"DICTAMEN FINAL:\s*\[(.*?)]", auditoria_resultado, re.DOTALL)
            auditoria_status = dictamen_final_match.group(1).strip() if dictamen_final_match else "❌ RECHAZADO (no se pudo extraer dictamen)"
            
            observaciones_start = auditoria_resultado.find("OBSERVACIONES FINALES:")
            audit_observations = auditoria_resultado[observaciones_start + len("OBSERVACIONES FINALES:"):].strip() if observaciones_start != -1 else "No se pudieron extraer observaciones."
            
            item_final_data = {
                "item_text": current_item_text, "classification": classification_details,
                "grafico_necesario": grafico_necesario, "descripcion_grafico": descripcion_grafico,
                "final_audit_status": auditoria_status, "final_audit_observations": audit_observations,
                "generation_prompt_used": full_generation_prompt, "auditor_prompt_used": full_auditor_prompt
            }

            if auditoria_status == "✅ CUMPLE TOTALMENTE":
                break
        
        except Exception as e:
            audit_observations = f"Error técnico durante la generación: {e}"
            auditoria_status = "❌ RECHAZADO (error técnico)"
            item_final_data = {
                "item_text": current_item_text if current_item_text else "No se pudo generar el ítem.",
                "classification": classification_details, "grafico_necesario": "NO", "descripcion_grafico": "",
                "final_audit_status": auditoria_status, "final_audit_observations": audit_observations,
                "generation_prompt_used": full_generation_prompt, "auditor_prompt_used": full_auditor_prompt
            }
            break

    return item_final_data

# --- Función para exportar preguntas a un documento Word ---
def exportar_a_word(preguntas_procesadas_list):
    doc = docx.Document()
    doc.add_heading('Preguntas Generadas y Auditadas', level=1)
    doc.add_paragraph('Este documento contiene los ítems generados por el sistema de IA y sus resultados de auditoría.\n')

    if not preguntas_procesadas_list:
        doc.add_paragraph('No se procesaron ítems para este informe.')

    for i, item_data in enumerate(preguntas_procesadas_list):
        doc.add_heading(f'Ítem #{i+1}', level=2)
        doc.add_paragraph('--- Clasificación del Ítem ---')
        for key, value in item_data["classification"].items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))
        
        doc.add_paragraph('\n' + item_data["item_text"])

        if item_data.get("grafico_necesario") == "SÍ" and item_data.get("descripcion_grafico"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("--- Gráfico Sugerido ---").bold = True
            doc.add_paragraph(f"**Tipo y Descripción del Gráfico:** {item_data['descripcion_grafico']}\n")

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("--- Resultado Final de Auditoría ---").bold = True
        doc.add_paragraph(f"**DICTAMEN FINAL:** {item_data.get('final_audit_status', 'N/A')}")
        doc.add_paragraph(f"**OBSERVACIONES FINALES:** {item_data.get('final_audit_observations', 'N/A')}\n")
        
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Interfaz de Usuario Principal de Streamlit ---
st.title("📚 Generador y Auditor de ítems para el proyecto SUMUN 🧠")
st.markdown("Esta aplicación genera ítems de selección múltiple y audita su calidad, usando la infraestructura de Vertex AI.")

# Sección de Carga de Archivos Global (Excel y PDF)
st.sidebar.header("Carga de Archivos Global")
uploaded_excel_file = st.sidebar.file_uploader("Sube tu archivo Excel (ESTRUCTURA_TOTAL.xlsx)", type=["xlsx"])
uploaded_pdf_file = st.sidebar.file_uploader("Sube tu archivo PDF (Manual_construccion_pruebas_IMProve.pdf)", type=["pdf"])

df_datos = None
manual_reglas_texto = ""

if uploaded_excel_file:
    df_datos = leer_excel_cargado(uploaded_excel_file)
if uploaded_pdf_file:
    manual_reglas_texto = leer_pdf_cargado(uploaded_pdf_file)
    max_manual_length = 15000
    if len(manual_reglas_texto) > max_manual_length:
        st.sidebar.warning(f"Manual truncado a {max_manual_length} caracteres.")
        manual_reglas_texto = manual_reglas_texto[:max_manual_length]
    st.sidebar.info(f"Manual de reglas cargado ({len(manual_reglas_texto)} caracteres).")

# --- Lógica principal de Generación y Auditoría de Ítems ---
st.header("Generación y Auditoría de Ítems")

if df_datos is None:
    st.info("Para comenzar, sube tu archivo Excel en la barra lateral.")
elif 'vertex_initialized' not in st.session_state or not st.session_state.vertex_initialized:
    st.warning("La conexión con Vertex AI no se ha establecido. Revisa la configuración del entorno.")
else:
    st.subheader("Selecciona los Criterios para la Generación")
    
    # Filtros de selección
    all_grades = df_datos['GRADO'].dropna().unique().tolist()
    grado_seleccionado = st.selectbox("Grado", sorted(all_grades))
    df_filtrado_grado = df_datos[df_datos['GRADO'].astype(str) == str(grado_seleccionado)]
    
    all_areas = df_filtrado_grado['ÁREA'].dropna().unique().tolist()
    area_seleccionada = st.selectbox("Área", sorted(all_areas))
    df_filtrado_area = df_filtrado_grado[df_filtrado_grado['ÁREA'] == area_seleccionada]
    
    all_asignaturas = df_filtrado_area['ASIGNATURA'].dropna().unique().tolist()
    asignatura_seleccionada = st.selectbox("Asignatura", sorted(all_asignaturas))
    df_filtrado_asignatura = df_filtrado_area[df_filtrado_area['ASIGNATURA'] == asignatura_seleccionada]
    
    all_estaciones = df_filtrado_asignatura['ESTACIÓN'].dropna().unique().tolist()
    estacion_seleccionada = st.selectbox("Estación", sorted(all_estaciones))
    df_filtrado_estacion = df_filtrado_asignatura[df_filtrado_asignatura['ESTACIÓN'] == estacion_seleccionada]
    
    st.subheader("Modo de Generación de Ítems")
    generate_all_for_station = st.checkbox("Generar TODOS los ítems de esta Estación")
    
    contexto_general_estacion = ""
    if generate_all_for_station:
        contexto_general_estacion = st.text_area("Escribe una idea para el contexto general de la estación (opcional, la IA puede crearlo):", height=150)

    proceso_cognitivo_seleccionado = None
    nanohabilidad_seleccionada = None
    df_item_seleccionado = None

    if not generate_all_for_station:
        all_procesos = df_filtrado_estacion['PROCESO COGNITIVO'].dropna().unique().tolist()
        proceso_cognitivo_seleccionado = st.selectbox("Proceso Cognitivo", sorted(all_procesos))
        df_filtrado_proceso = df_filtrado_estacion[df_filtrado_estacion['PROCESO COGNITIVO'] == proceso_cognitivo_seleccionado]
        
        all_nanohabilidades = df_filtrado_proceso['NANOHABILIDAD'].dropna().unique().tolist()
        nanohabilidad_seleccionada = st.selectbox("Nanohabilidad", sorted(all_nanohabilidades))
        df_item_seleccionado = df_filtrado_proceso[df_filtrado_proceso['NANOHABILIDAD'] == nanohabilidad_seleccionada]
    else:
        df_item_seleccionado = df_filtrado_estacion.copy()

    if df_item_seleccionado.empty:
        st.error("No hay datos para generar con los filtros actuales.")
    else:
        informacion_adicional_usuario = ""
        if not generate_all_for_station:
            informacion_adicional_usuario = st.text_area("Contexto adicional para este ítem individual (Opcional):")

        st.subheader("Personaliza con Prompts Adicionales (Opcional)")
        use_additional_prompts = st.checkbox("Activar Prompts Adicionales")
        prompt_bloom_adicional, prompt_construccion_adicional, prompt_especifico_adicional, prompt_auditor_adicional = "", "", "", ""

        if use_additional_prompts:
            # ... (código para los prompts adicionales, se mantiene igual)

        st.subheader("Configuración de Modelos de IA (Vertex AI)")
        col1, col2 = st.columns(2)
        with col1:
            gen_model_name = st.selectbox("Modelo para Generación", ["gemini-2.0-flash", "gemini-2.5-pro", "gemini-2.0-flash-lite"])
        with col2:
            audit_model_name = st.selectbox("Modelo para Auditoría", ["gemini-2.0-flash-lite", "gemini-2.0-flash", "gemini-2.5-pro"])

        if st.button("Generar y Auditar Ítem(s)"):
            criterios_para_preguntas = {
                "tipo_pregunta": "opción múltiple con 4 opciones", "dificultad": "media",
                "contexto_educativo": "estudiantes de preparatoria (bachillerato)",
                "formato_justificacion": """
                    • Justificación correcta: debe explicar el razonamiento (NO por descarte).
                    • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
                """
            }
            processed_items_list = []

            if generate_all_for_station:
                st.info(f"Generando ítems para la Estación: {estacion_seleccionada}")
                unique_procesos = df_item_seleccionado[['PROCESO COGNITIVO', 'NANOHABILIDAD', 'MICROHABILIDAD', 'COMPETENCIA NANOHABILIDAD']].drop_duplicates().to_dict('records')
                progress_bar = st.progress(0)
                
                for i, item_spec_row in enumerate(unique_procesos):
                    st.write(f"Procesando {i+1}/{len(unique_procesos)}: {item_spec_row['PROCESO COGNITIVO']}...")
                    current_fila_datos = {'GRADO': grado_seleccionado, 'ÁREA': area_seleccionada, 'ASIGNATURA': asignatura_seleccionada, 'ESTACIÓN': estacion_seleccionada, **item_spec_row}
                    item_data = generar_pregunta_con_seleccion(gen_model_name, audit_model_name, current_fila_datos, criterios_para_preguntas, manual_reglas_texto, informacion_adicional_usuario, prompt_bloom_adicional, prompt_construccion_adicional, prompt_especifico_adicional, prompt_auditor_adicional, contexto_general_estacion)
                    if item_data:
                        processed_items_list.append(item_data)
                    progress_bar.progress((i + 1) / len(unique_procesos))
                st.success("Proceso completado.")

            else:
                st.info(f"Generando ítem individual para: {nanohabilidad_seleccionada}")
                item_data = generar_pregunta_con_seleccion(gen_model_name, audit_model_name, df_item_seleccionado.iloc[0], criterios_para_preguntas, manual_reglas_texto, informacion_adicional_usuario, prompt_bloom_adicional, prompt_construccion_adicional, prompt_especifico_adicional, prompt_auditor_adicional)
                if item_data:
                    processed_items_list.append(item_data)
                    st.success(f"Ítem generado. Dictamen: {item_data.get('final_audit_status')}")

            st.session_state['processed_items_list'] = processed_items_list
            
            if processed_items_list:
                st.subheader("Resumen del Primer Ítem Procesado:")
                first_item = processed_items_list[0]
                st.markdown(first_item['item_text'])
                st.info(f"Dictamen: {first_item['final_audit_status']}")
            else:
                st.error("No se pudo generar ningún ítem.")
        
        st.header("Exportar Resultados")
        if 'processed_items_list' in st.session_state and st.session_state['processed_items_list']:
            num_items = len(st.session_state['processed_items_list'])
            st.write(f"Hay {num_items} ítem(s) listos para exportar.")
            
            word_buffer = exportar_a_word(st.session_state['processed_items_list'])
            st.download_button(
                label="Descargar Ítem(s) en Word", data=word_buffer,
                file_name=f"items_{estacion_seleccionada.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            combined_prompts_content = ""
            for i, item_data in enumerate(st.session_state['processed_items_list']):
                combined_prompts_content += f"--- ÍTEM #{i+1} ({item_data['classification']['Proceso Cognitivo']}) ---\n"
                combined_prompts_content += f"--- PROMPT GENERADOR ---\n{item_data.get('generation_prompt_used', 'N/A')}\n\n"
                combined_prompts_content += f"--- PROMPT AUDITOR ---\n{item_data.get('auditor_prompt_used', 'N/A')}\n\n{'='*80}\n\n"

            st.download_button(
                label="Descargar Prompts como TXT", data=combined_prompts_content.encode('utf-8'),
                file_name=f"prompts_{estacion_seleccionada.replace(' ', '_')}.txt", mime="text/plain"
            )
